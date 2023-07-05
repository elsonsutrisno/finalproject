from datetime import time
import openpyxl
from django.http import HttpResponse
import os
from .models import table_session, table_time

def fetch_all_session_objects():
    return table_session.objects.all()

def fetch_time_objects(session_id):
    session = table_session.objects.get(id=session_id)
    return session.table_time_set.all()

def get_session_by_id(session_id):
    return table_session.objects.get(pk=session_id)

def get_time_by_id(time_id):
    return table_time.objects.get(pk=time_id)

def update_session_limit(form_session):
    form_session.save()

def update_session_available_seat(time_id, limit_after):
    time_object =  table_time.objects.get(pk=time_id)
    limit_before = time_object.seat_limit
    available_before = time_object.available_seat
    difference = abs(limit_before - limit_after)
    print(limit_before, limit_after, difference, available_before)
    print(time_object, type(time_object))

    if available_before < time_object.available_seat:
        return False

    if limit_before > limit_after: 
        time_object.available_seat = available_before - difference
        print(f"Lebih besar {limit_before, limit_after}")
    elif limit_before < limit_after: 
        value = (available_before + difference)
        time_object.available_seat = value
        print(f"Lebih kecil = {available_before} + {difference} = {time_object.available_seat}")
    else: 
        time_object.available_seat = time_object.seat_limit
    
    time_object.seat_limit = limit_after
    time_object.save()
    return True

def delete_session_object(session_id):
    session_id.delete()

def delete_time_object(time_id):
    time_id.delete()

def export_data_to_excel(file_path):
    workbook = openpyxl.Workbook()
    worksheet = workbook.active

    headers = ["Date", "Name", "Time", "Available", "Limit", "Menu"]
    for col_num, header in enumerate(headers, 1):
        worksheet.cell(row=1, column=col_num, value=header)

    sessions = fetch_all_session_objects()
    for session in sessions:
        times = fetch_time_objects(session.id)

        for time_obj in times:
            date = session.date
            name = session.get_name_display()
            menu = session.menu

            time_value = time_obj.time
            available = time_obj.available_seat
            limit = time_obj.seat_limit

            row_data = [date, name, time_value, available, limit, menu]
            row_num = worksheet.max_row + 1
            for col_num, value in enumerate(row_data, 1):
                worksheet.cell(row=row_num, column=col_num, value=value)

    workbook.save(file_path)

def save_session_and_times(form_session, form_time):
    session = form_session.save()

    if session.name == "Breakfast":
        times = [
            time(7, 0),
            time(7, 30),
            time(8, 0),
            time(8, 30),
            time(9, 0),
        ]
    elif session.name == "Lunch":
        times = [
            time(11, 0),
            time(11, 30),
            time(12, 0),
            time(12, 30),
            time(13, 0),
            time(13, 30),
        ]
    elif session.name == "Dinner":
        times = [
            time(17, 0),
            time(17, 30),
            time(18, 0),
            time(18, 30),
            time(19, 0),
            time(19, 30),
        ]
    else:
        times = []

    for time_value in times:
        table_time.objects.create(
            time=time_value,
            session_id=session,
            seat_limit=form_time.cleaned_data["seat_limit"],
            available_seat = form_time.cleaned_data["seat_limit"],
        )

def download_file_response(file_path):
    with open(file_path, 'rb') as file:
        file_data = file.read()

    response = HttpResponse(file_data, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="order_record.xlsx"'

    os.remove(file_path)
    return response

# DOWNLOAD REPORT FUNTIONS START #
import os
from django.http import HttpResponse
from datetime import datetime
from docx import Document
from io import BytesIO

def fetch_sessions_by_date_range(start_date, end_date):
    sessions = table_session.objects.filter(date__range=[start_date, end_date])
    return sessions
    
def download_report_doc(start_date, end_date, filename):
    # Fetch session objects within the specified date range
    sessions = table_session.objects.filter(date__range=[start_date, end_date])

    # Calculate total available seats and seating capacity limit
    total_available_seats = 0
    seating_capacity_limit = 0

    for session in sessions:
        times = table_time.objects.filter(session_id=session)

        for time_obj in times:
            total_available_seats += time_obj.available_seat or time_obj.seat_limit
            seating_capacity_limit += time_obj.seat_limit

    # Calculate the percentage of seats occupied and seating capacity utilized
    total_occupied_seats = seating_capacity_limit - total_available_seats
    percentage_occupied = (total_occupied_seats / seating_capacity_limit) * 100
    percentage_utilized = (total_available_seats / seating_capacity_limit) * 100

    # Generate the Word document in memory
    document = Document()
    document.add_heading('Dining Hall Booking Report', level=1)

    # Add report summary section
    document.add_paragraph(f"Date Range: {start_date} to {end_date}")
    report_summary_paragraph = document.add_paragraph()
    report_summary_run = report_summary_paragraph.add_run("Report Summary:")
    report_summary_run.bold = True
    document.add_paragraph(f"Available Seats: {total_available_seats}")
    document.add_paragraph(f"Seats Occupied: {percentage_occupied}%")
    document.add_paragraph(f"Seating Capacity Limit: {seating_capacity_limit}")
    document.add_paragraph(f"Capacity Utilization: {percentage_utilized}%")

    # Add booking status analysis section
    analysis_paragraph = document.add_paragraph()
    analysis_run = analysis_paragraph.add_run("Analysis:")
    analysis_run.bold = True
    document.add_paragraph(f"The number of available seats during the specified time range is {total_available_seats}.")
    document.add_paragraph(f"The percentage of seats occupied within the specified time range is {percentage_occupied}%.")
    document.add_paragraph(f"The maximum seating capacity allowed at Dining Hall is {seating_capacity_limit}.")
    document.add_paragraph(f"The percentage of seating capacity utilized within the specified time range is {percentage_utilized}%.")

    # Add recommendations section
    report_summary_paragraph = document.add_paragraph()
    report_summary_run = report_summary_paragraph.add_run("Report Summary:")
    report_summary_run.bold = True
    document.add_paragraph("If the percentage of seats occupied is consistently high, consider increasing seating capacity or optimizing the booking system to accommodate more guests.")
    document.add_paragraph("If the percentage of seating capacity utilized exceeds the limit, take measures to ensure compliance, such as managing bookings or implementing a waitlist system.")

    # Add contact information
    document.add_paragraph(f"For further assistance, please contact [Restaurant Contact Number] or visit [Restaurant Website/Email Address].")
    document.add_paragraph("Thank you for choosing Dining Hall.")

    document.add_page_break()  # Add a page break
    document.add_heading("Booking Report's Table", level=1)

    headers = ["Date", "Name", "Time Range", "Available", "Limit", "Menu"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    total_available = 0
    total_limit = 0

    for session in sessions:
        times = fetch_time_objects(session.id)

        # Calculate start and end times for this session
        start_time = min(time_obj.time for time_obj in times)
        end_time = max(time_obj.time for time_obj in times)
        time_range = f"{start_time} - {end_time}"

        # Calculate available and limit seats for this session
        available = sum(time_obj.available_seat or time_obj.seat_limit for time_obj in times)
        limit = sum(time_obj.seat_limit for time_obj in times)

        # Add row to table
        row = table.add_row().cells
        row[0].text = str(session.date)
        row[1].text = session.get_name_display()
        row[2].text = time_range
        row[3].text = str(available)
        row[4].text = str(limit)
        row[5].text = session.menu

        # Update the total counts
        total_available += available
        total_limit += limit

    # Add a summary row for the total count
    total_row = table.add_row().cells
    total_row[0].text = "Total"
    total_row[3].text = str(total_available)
    total_row[4].text = str(total_limit)

    # Save the Word document to a BytesIO object
    file_data = BytesIO()
    document.save(file_data)

    # Set the response headers and content
    response = HttpResponse(file_data.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Booking_Report from {start_date} to {end_date}.docx"'

    return response

def validate_dates(start_date: str, end_date: str, date_format: str = '%Y-%m-%d') -> bool:
    """
    Validate if start_date and end_date are valid dates in the specified format.

    :param start_date: The start date string to validate.
    :param end_date: The end date string to validate.
    :param date_format: The date format to use for validation (default: '%Y-%m-%d').
    :return: True if both start_date and end_date are valid dates in the specified format, False otherwise.
    """
    try:
        parsed_start_date = datetime.strptime(start_date, date_format).date()
        parsed_end_date = datetime.strptime(end_date, date_format).date()
        return True, parsed_start_date, parsed_end_date
    except ValueError:
        return False
    
# DOWNLOAD REPORT FUNTIONS END #